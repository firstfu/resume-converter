import os
import shutil
import tempfile
from datetime import datetime
from pathlib import Path

import pytesseract
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from fastapi import FastAPI, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from PIL import Image

app = FastAPI(title="Resume Converter API")

# 配置 CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 創建臨時文件夾
UPLOAD_DIR = Path("temp_uploads")
UPLOAD_DIR.mkdir(exist_ok=True)


def process_image_with_tesseract(image_path: str) -> str:
    """
    使用 Tesseract 處理圖像
    """
    try:
        return pytesseract.image_to_string(Image.open(image_path), lang="eng+chi_tra")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Tesseract處理失敗: {str(e)}")


def create_word_document(text: str) -> Path:
    """
    將文本轉換為 Word 文檔
    """
    try:
        # 創建新的 Word 文檔
        doc = Document()

        # 設置基本樣式
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Arial"
        font.size = Pt(12)

        # 添加標題
        heading = doc.add_heading("簡歷內容", 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 處理文本內容
        paragraphs = text.split("\n\n")
        for para in paragraphs:
            if para.strip():
                p = doc.add_paragraph()
                p.add_run(para.strip())

        # 保存文檔
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = UPLOAD_DIR / f"resume_{timestamp}.docx"
        doc.save(str(output_path))

        return output_path
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Word文檔生成失敗: {str(e)}")


@app.post("/api/upload")
async def upload_file(file: UploadFile):
    """
    處理上傳的簡歷圖像文件
    """
    # 檢查文件類型
    allowed_types = {"image/jpeg", "image/png", "application/pdf"}
    if file.content_type not in allowed_types:
        raise HTTPException(status_code=400, detail="不支持的文件類型")

    # 保存上傳的文件
    temp_file = Path(tempfile.gettempdir()) / f"upload_{file.filename}"
    try:
        with temp_file.open("wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        # 處理圖像
        text_result = process_image_with_tesseract(str(temp_file))

        # 生成 Word 文檔
        docx_path = create_word_document(text_result)

        return JSONResponse(
            {"status": "success", "text": text_result, "docx_file": docx_path.name}
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

    finally:
        # 清理臨時文件
        if temp_file.exists():
            temp_file.unlink()


@app.get("/api/download/{filename}")
async def download_file(filename: str):
    """
    下載生成的 Word 文檔
    """
    file_path = UPLOAD_DIR / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="文件不存在")

    return FileResponse(
        path=file_path,
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# 掛載靜態文件
app.mount("/", StaticFiles(directory="static", html=True), name="static")

if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=8000, reload=True)
